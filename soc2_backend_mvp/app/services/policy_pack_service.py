import io
import zipfile
from datetime import datetime
from docx import Document

from app.services.gemini_service import gemini_generate_text


POLICIES = [
    ("Access_Control_Policy.docx", "Access Control Policy"),
    ("Incident_Response_Plan.docx", "Incident Response Plan"),
    ("Change_Management_Policy.docx", "Change Management Policy"),
    ("Vendor_Risk_Management_Policy.docx", "Vendor Risk Management Policy"),
    ("Data_Retention_Policy.docx", "Data Retention Policy"),
]


def build_policy_prompt(company_profile: dict, policy_title: str) -> str:
    company_name = company_profile.get("company", {}).get("legalName", "Company")
    industry = company_profile.get("company", {}).get("industry", "SaaS")
    cloud = company_profile.get("itEnvironment", {}).get("cloudProvider", "AWS")

    return f"""
You are a SOC2 compliance consultant.

Write a SOC2-ready policy document.

Company name: {company_name}
Industry: {industry}
Cloud provider: {cloud}

Policy title: {policy_title}

Rules:
- Use clear headings
- Include: Purpose, Scope, Roles & Responsibilities, Policy, Procedures, Evidence, Review Frequency
- Make it audit-friendly
- Keep it professional
""".strip()


def fallback_policy_text(title: str) -> str:
    return f"""
{title}

Purpose
This policy defines the organization's approach to {title.lower()}.

Scope
Applies to all employees, contractors, and systems.

Roles & Responsibilities
- Security Team: owns the policy
- IT Admins: enforce controls
- Employees: follow the policy

Policy
- Controls must be enforced consistently
- Exceptions require approval and documentation

Evidence
- Access logs
- Ticketing records
- Review checklists

Review Frequency
Quarterly
""".strip()


def text_to_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()

    # Title
    doc.add_heading(title, level=1)

    # Split into paragraphs
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        # basic heading detection
        if line.lower() in ["purpose", "scope", "roles & responsibilities", "policy", "procedures", "evidence", "review frequency"]:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def generate_policy_pack_zip(company_profile: dict) -> bytes:
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, title in POLICIES:
            prompt = build_policy_prompt(company_profile, title)

            # Gemini optional (fallback if quota issues)
            content = gemini_generate_text(prompt)
            if not content:
                content = fallback_policy_text(title)

            docx_bytes = text_to_docx_bytes(title, content)
            zf.writestr(filename, docx_bytes)

        meta = f"GeneratedAt: {datetime.utcnow().isoformat()}Z\n"
        zf.writestr("meta.txt", meta)

    buffer.seek(0)
    return buffer.read()
