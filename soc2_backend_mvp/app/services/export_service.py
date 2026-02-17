import io
import zipfile
from bson import ObjectId
from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup

from ..db import get_db
from ..utils.errors import not_found


def _markdown_to_text(md: str) -> str:
    html = markdown(md)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n")


async def export_docx_single(user_id: str, doc_id: str) -> tuple[bytes, str]:
    db = get_db()
    doc = await db.documents.find_one({"_id": ObjectId(doc_id), "userId": user_id})
    if not doc:
        not_found("Document not found")

    d = Document()
    d.add_heading(doc["title"], level=1)

    text = _markdown_to_text(doc["contentMarkdown"])
    for line in text.splitlines():
        if line.strip():
            d.add_paragraph(line)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue(), doc["title"]


async def export_docx_pack_zip(user_id: str) -> bytes:
    db = get_db()
    docs = []
    cursor = db.documents.find({"userId": user_id}).sort("title", 1)
    async for x in cursor:
        docs.append(x)

    if not docs:
        not_found("No documents to export")

    zip_buf = io.BytesIO()

    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
        for doc in docs:
            d = Document()
            d.add_heading(doc["title"], level=1)
            text = _markdown_to_text(doc["contentMarkdown"])
            for line in text.splitlines():
                if line.strip():
                    d.add_paragraph(line)

            file_buf = io.BytesIO()
            d.save(file_buf)

            safe_name = doc["title"].replace("/", "-").replace("\\", "-")
            z.writestr(f"{safe_name}.docx", file_buf.getvalue())

    return zip_buf.getvalue()
